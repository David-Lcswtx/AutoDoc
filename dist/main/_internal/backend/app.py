from flask import Flask, request, jsonify, send_from_directory
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog
import threading
from datetime import datetime

app = Flask(__name__, static_folder="../frontend", static_url_path="")


def abrir_seletor(tipo_seletor):
    resultado = []

    def target():
        root = tk.Tk()
        root.withdraw()
        root.lift()
        root.attributes("-topmost", True)

        if tipo_seletor == "modelo":
            caminho = filedialog.askopenfilename(
                title="Selecionar Modelo .docx", filetypes=[("Arquivos Word", "*.docx")]
            )
        elif tipo_seletor == "destino":
            caminho = filedialog.askdirectory(
                title="Escolha a Pasta onde o arquivo será gerado"
            )

        resultado.append(caminho)
        root.destroy()

    thread = threading.Thread(target=target)
    thread.start()
    thread.join()

    return resultado[0] if resultado else ""


@app.route("/")
def index():
    return send_from_directory(app.static_folder, "index.html")


@app.route("/selecionar-modelo", methods=["GET"])
def selecionar_modelo():
    try:
        caminho_arquivo = abrir_seletor("modelo")
        return jsonify({"caminho": caminho_arquivo}), 200
    except Exception as e:
        return jsonify({"erro": str(e)}), 500


@app.route("/selecionar-destino", methods=["GET"])
def selecionar_destino():
    try:
        caminho_pasta = abrir_seletor("destino")
        return jsonify({"caminho": caminho_pasta}), 200
    except Exception as e:
        return jsonify({"erro": str(e)}), 500


def substituir_tags(doc, dados_mapa):
    for paragrafo in doc.paragraphs:
        for tag, valor in dados_mapa.items():
            if tag in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(tag, str(valor))

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for tag, valor in dados_mapa.items():
                        if tag in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(tag, str(valor))


@app.route("/gerar-documento", methods=["POST"])
def gerar_documento():
    dados = request.json

    try:
        caminho_modelo = dados["config"]["modelo"]
        pasta_destino = dados["config"][
            "destino"
        ]  # Agora isso é apenas a pasta selecionada
        nome_cliente = dados["cliente"]["nome"]

        if not caminho_modelo or not os.path.exists(caminho_modelo):
            caminho_modelo = os.path.join("documentos", "modelo.docx")
            pasta_destino = "documentos"

        if not os.path.exists(caminho_modelo):
            return jsonify({"mensagem": "Arquivo de modelo não encontrado."}), 400

        doc = Document(caminho_modelo)

        # Mapeamento do Cliente e Endereço
        mapeamento = {
            "{{CLIENTE_NOME}}": nome_cliente,
            "{{CLIENTE_NACIONALIDADE}}": dados["cliente"]["nacionalidade"],
            "{{CLIENTE_ESTADO_CIVIL}}": dados["cliente"]["estado_civil"],
            "{{CLIENTE_DATA_NASC}}": dados["cliente"]["data_nascimento"],
            "{{CLIENTE_CPF}}": dados["cliente"]["cpf"],
            "{{CLIENTE_RG}}": dados["cliente"]["rg"],
            "{{END_CEP}}": dados["endereco"]["cep"],
            "{{END_LOGRADOURO}}": dados["endereco"]["logradouro"],
            "{{END_NUMERO}}": dados["endereco"]["numero"],
            "{{END_COMPLEMENTO}}": dados["endereco"]["complemento"],
            "{{END_BAIRRO}}": dados["endereco"]["bairro"],
            "{{END_MUNICIPIO}}": dados["endereco"]["municipio"],
            "{{END_UF}}": dados["endereco"]["uf"],
        }

        if dados["possui_representante"] and dados["representante"]:
            rep = dados["representante"]
            mapeamento.update(
                {
                    "{{REP_NOME}}": rep["nome"],
                    "{{REP_NACIONALIDADE}}": rep["nacionalidade"],
                    "{{REP_ESTADO_CIVIL}}": rep["estado_civil"],
                    "{{REP_CPF}}": rep["cpf"],
                    "{{REP_RG}}": rep["rg"],
                }
            )
        else:
            mapeamento.update(
                {
                    "{{REP_NOME}}": "",
                    "{{REP_NACIONALIDADE}}": "",
                    "{{REP_ESTADO_CIVIL}}": "",
                    "{{REP_CPF}}": "",
                    "{{REP_RG}}": "",
                }
            )

        substituir_tags(doc, mapeamento)

        # Garante que a pasta destino existe
        os.makedirs(pasta_destino, exist_ok=True)

        # Cria um nome automático bem organizado para a empresa, ex: "Processo_David_Lucas.docx"
        # Se o campo de nome estiver vazio, ele usa "Cliente_Sem_Nome"
        nome_limpo = (
            nome_cliente.strip().replace(" ", "_")
            if nome_cliente.strip()
            else "Cliente_Sem_Nome"
        )
        nome_arquivo = f"Processo_{nome_limpo}.docx"

        # Junta a pasta selecionada com o nome do arquivo gerado automaticamente
        caminho_final_salvamento = os.path.join(pasta_destino, nome_arquivo)

        doc.save(caminho_final_salvamento)

        return (
            jsonify(
                {
                    "mensagem": f"Documento salvo automaticamente como '{nome_arquivo}' na pasta selecionada!"
                }
            ),
            200,
        )

    except Exception as e:
        return (
            jsonify({"mensagem": f"Erro interno ao gerar o documento: {str(e)}"}),
            500,
        )


if __name__ == "__main__":
    app.run(debug=True)
